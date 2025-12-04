"""
Генератор Word документов
"""

import os
from typing import Dict, Any
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
from html.parser import HTMLParser


class WordGenerator:
    """Генератор документов в формате Word (.docx)"""
    
    def generate(self, template_path: str, data: Dict[str, Any], output_path: str) -> str:
        """
        Генерация Word документа из шаблона
        
        Args:
            template_path: Путь к шаблону Word
            data: Словарь с данными для подстановки
            output_path: Путь для сохранения результата
            
        Returns:
            Путь к сгенерированному файлу
        """
        if not os.path.exists(template_path):
            # Создаем новый документ, если шаблон не существует
            doc = Document()
            doc.add_heading('Документ', 0)
            doc.add_paragraph('{{content}}')
        else:
            doc = Document(template_path)
        
        # Отладочный вывод
        print(f"Замена переменных в документе. Доступные данные: {list(data.keys())}")
        print(f"Пример данных: contract_number={data.get('contract_number')}, date={data.get('date')}")
        print(f"Content HTML: {data.get('content_html', 'НЕТ')}")
        
        # Заменяем переменные в документе
        self._replace_variables(doc, data)
        
        # Добавляем дополнительное содержимое из WYSIWYG редактора, если есть
        html_content = data.get('content_html', '').strip() if data.get('content_html') else ''
        if html_content and html_content not in ['<p><br></p>', '<p></p>', '']:
            print(f"Добавление HTML контента в документ. Длина: {len(html_content)}")
            print(f"HTML контент (первые 200 символов): {html_content[:200]}")
            
            # Ищем место для вставки - ищем переменную {{content}} или добавляем в конец
            content_added = False
            for paragraph in doc.paragraphs:
                para_text = paragraph.text
                if '{{content}}' in para_text or '{content}' in para_text:
                    # Заменяем переменную content на HTML контент
                    print("Найдена переменная {{content}}, заменяем...")
                    self._add_html_content_to_paragraph(paragraph, html_content)
                    content_added = True
                    break
            
            # Если переменной {{content}} не найдено, добавляем в конец документа
            if not content_added:
                print("Переменная {{content}} не найдена, добавляем в конец документа...")
                doc.add_paragraph()  # Пустая строка
                doc.add_heading('Дополнительное содержимое', level=2)
                self._add_html_content(doc, html_content)
        
        # Сохраняем документ
        os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else '.', exist_ok=True)
        doc.save(output_path)
        
        return output_path
    
    def _replace_variables(self, doc: Document, data: Dict[str, Any]):
        """
        Замена переменных в документе
        
        Args:
            doc: Объект документа Word
            data: Словарь с данными
        """
        # Обрабатываем параграфы
        for paragraph in doc.paragraphs:
            self._replace_in_paragraph(paragraph, data)
        
        # Обрабатываем таблицы
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, data)
        
        # Обрабатываем заголовки (headings) - они тоже могут содержать переменные
        # Заголовки уже обрабатываются в цикле выше, но убедимся что все обработано
    
    def _replace_in_paragraph(self, paragraph, data: Dict[str, Any]):
        """
        Замена переменных в параграфе
        
        Args:
            paragraph: Параграф документа
            data: Словарь с данными
        """
        text = paragraph.text
        original_text = text
        
        if not text or ('{{' not in text and '{' not in text):
            # Нет переменных для замены
            return
        
        # Ищем все переменные в формате {{variable_name}} или {variable_name}
        # Поддерживаем оба формата для совместимости
        # Сначала заменяем двойные скобки {{variable}}, чтобы не конфликтовать с одинарными
        pattern_double = r'\{\{(\w+)\}\}'  # {{variable}}
        
        def replace_double(match):
            var_name = match.group(1)
            value = data.get(var_name, '')
            result = str(value) if value else ''
            print(f"  Замена {{{{{var_name}}}}}: '{result}'")
            return result
        
        text = re.sub(pattern_double, replace_double, text)
        
        # Потом заменяем одинарные скобки {variable}, но только те, что не являются частью {{variable}}
        # Используем негативный lookahead и lookbehind
        pattern_single = r'(?<!\{)\{(\w+)\}(?!\})'  # {variable} но не часть {{variable}}
        
        def replace_single(match):
            var_name = match.group(1)
            value = data.get(var_name, '')
            result = str(value) if value else ''
            print(f"  Замена {{{var_name}}}: '{result}'")
            return result
        
        text = re.sub(pattern_single, replace_single, text)
        
        # Если текст изменился, обновляем параграф
        if text != original_text:
            # Сохраняем форматирование
            runs = paragraph.runs
            paragraph.clear()
            
            # Добавляем новый текст с сохранением базового форматирования
            run = paragraph.add_run(text)
            if runs:
                run.font.size = runs[0].font.size if runs[0].font.size else Pt(11)
                run.font.name = runs[0].font.name if runs[0].font.name else 'Times New Roman'
    
    def create_from_scratch(self, data: Dict[str, Any], output_path: str) -> str:
        """
        Создание Word документа с нуля
        
        Args:
            data: Данные для документа
            output_path: Путь для сохранения
            
        Returns:
            Путь к созданному файлу
        """
        doc = Document()
        
        # Заголовок
        if 'title' in data:
            doc.add_heading(data['title'], 0)
        
        # Содержание
        if 'content_html' in data:
            # Используем HTML контент из WYSIWYG редактора
            self._add_html_content(doc, data['content_html'])
        elif 'content' in data:
            if isinstance(data['content'], list):
                for item in data['content']:
                    doc.add_paragraph(item, style='List Bullet')
            else:
                doc.add_paragraph(str(data['content']))
        
        # Таблица, если есть
        if 'table_data' in data:
            table = doc.add_table(rows=1, cols=len(data['table_data'][0]))
            table.style = 'Light Grid Accent 1'
            
            # Заголовки
            header_cells = table.rows[0].cells
            for i, header in enumerate(data['table_data'][0]):
                header_cells[i].text = str(header)
            
            # Данные
            for row_data in data['table_data'][1:]:
                row_cells = table.add_row().cells
                for i, cell_data in enumerate(row_data):
                    row_cells[i].text = str(cell_data)
        
        os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else '.', exist_ok=True)
        doc.save(output_path)
        
        return output_path
    
    def _add_html_content(self, doc: Document, html_content: str):
        """
        Добавление HTML контента в Word документ
        
        Args:
            doc: Объект документа Word
            html_content: HTML контент из WYSIWYG редактора
        """
        if not html_content or html_content.strip() == '' or html_content.strip() in ['<p><br></p>', '<p></p>']:
            return
        
        import html as html_module
        
        # Парсим HTML параграфы из Quill (формат <p>...</p>)
        para_pattern = r'<p[^>]*>(.*?)</p>'
        paragraphs_html = re.findall(para_pattern, html_content, re.DOTALL | re.IGNORECASE)
        
        if not paragraphs_html:
            # Если нет параграфов, извлекаем весь текст
            text = re.sub(r'<[^>]+>', '', html_content)
            text = html_module.unescape(text)
            text = text.replace('&nbsp;', ' ')
            if text.strip():
                para = doc.add_paragraph()
                self._add_formatted_run(para, text.strip(), html_content)
        else:
            for para_html in paragraphs_html:
                para_html = para_html.strip()
                if not para_html or para_html == '<br>':
                    doc.add_paragraph()  # Пустой параграф
                    continue
                
                para = doc.add_paragraph()
                text = re.sub(r'<[^>]+>', '', para_html)
                text = html_module.unescape(text)
                text = text.replace('&nbsp;', ' ')
                
                if text.strip():
                    self._add_formatted_run(para, text.strip(), para_html)
    
    def _add_formatted_run(self, paragraph, text: str, html_content: str):
        """
        Добавление форматированного текста в параграф
        
        Args:
            paragraph: Параграф документа
            text: Текстовое содержимое
            html_content: HTML для определения форматирования
        """
        # Проверяем наличие форматирования
        is_bold = '<strong>' in html_content or '<b>' in html_content
        is_italic = '<em>' in html_content or '<i>' in html_content
        is_underline = '<u>' in html_content
        
        run = paragraph.add_run(text)
        
        if is_bold:
            run.bold = True
        if is_italic:
            run.italic = True
        if is_underline:
            run.underline = True
    
    def _add_html_content_to_paragraph(self, paragraph, html_content: str):
        """
        Замена содержимого параграфа на HTML контент
        
        Args:
            paragraph: Параграф документа
            html_content: HTML контент
        """
        if not html_content or html_content.strip() == '' or html_content.strip() in ['<p><br></p>', '<p></p>']:
            paragraph.clear()
            return
        
        # Очищаем параграф
        paragraph.clear()
        
        # Парсим HTML и добавляем в параграф
        import html as html_module
        
        # Парсим HTML параграфы из Quill
        para_pattern = r'<p[^>]*>(.*?)</p>'
        paragraphs_html = re.findall(para_pattern, html_content, re.DOTALL | re.IGNORECASE)
        
        if paragraphs_html:
            # Если есть параграфы, берем первый (для замены в одном параграфе)
            para_html = paragraphs_html[0].strip()
            if para_html and para_html != '<br>':
                text = re.sub(r'<[^>]+>', '', para_html)
                text = html_module.unescape(text)
                text = text.replace('&nbsp;', ' ')
                
                if text.strip():
                    self._add_formatted_run(paragraph, text.strip(), para_html)
        else:
            # Если нет параграфов, извлекаем весь текст
            text = re.sub(r'<[^>]+>', '', html_content)
            text = html_module.unescape(text)
            text = text.replace('&nbsp;', ' ')
            
            if text.strip():
                self._add_formatted_run(paragraph, text.strip(), html_content)

