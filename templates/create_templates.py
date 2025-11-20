"""
Скрипт для создания шаблонов документов
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Создаем директорию для шаблонов
os.makedirs('templates', exist_ok=True)

# Шаблон договора
def create_contract_template():
    doc = Document()
    
    # Заголовок
    title = doc.add_heading('ДОГОВОР', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Номер договора
    doc.add_paragraph('№ {{contract_number}}')
    
    # Дата
    doc.add_paragraph(f'Дата: {{date}}')
    doc.add_paragraph('')
    
    # Стороны
    doc.add_paragraph('{{party1_name}}, именуемое в дальнейшем "Заказчик",')
    doc.add_paragraph('и {{party2_name}}, именуемое в дальнейшем "Исполнитель",')
    doc.add_paragraph('заключили настоящий договор о нижеследующем:')
    doc.add_paragraph('')
    
    # Предмет договора
    doc.add_heading('1. ПРЕДМЕТ ДОГОВОРА', level=1)
    doc.add_paragraph('{{subject}}')
    doc.add_paragraph('')
    
    # Стоимость
    doc.add_heading('2. СТОИМОСТЬ И ПОРЯДОК РАСЧЕТОВ', level=1)
    doc.add_paragraph('Стоимость работ составляет {{amount}} рублей.')
    doc.add_paragraph('')
    
    # Сроки
    doc.add_heading('3. СРОКИ ВЫПОЛНЕНИЯ', level=1)
    doc.add_paragraph('Работы должны быть выполнены до {{deadline}}.')
    doc.add_paragraph('')
    
    # Подписи
    doc.add_paragraph('ЗАКАЗЧИК:')
    doc.add_paragraph('_________________ {{customer_signature}}')
    doc.add_paragraph('')
    doc.add_paragraph('ИСПОЛНИТЕЛЬ:')
    doc.add_paragraph('_________________ {{executor_signature}}')
    
    doc.save('templates/contract_template.docx')
    print("Шаблон договора создан: templates/contract_template.docx")

# Шаблон финансового отчета
def create_report_template():
    doc = Document()
    
    # Заголовок
    title = doc.add_heading('ФИНАНСОВЫЙ ОТЧЕТ', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Организация
    doc.add_paragraph('Организация: {{company_name}}')
    doc.add_paragraph('ИНН: {{inn}}')
    doc.add_paragraph('ОГРН: {{ogrn}}')
    doc.add_paragraph('')
    
    # Период
    doc.add_paragraph('Период: {{period_start}} - {{period_end}}')
    doc.add_paragraph('Дата формирования: {{date}}')
    doc.add_paragraph('')
    
    # Основные показатели
    doc.add_heading('Основные показатели', level=1)
    doc.add_paragraph('Выручка: {{revenue}} руб.')
    doc.add_paragraph('Расходы: {{expenses}} руб.')
    doc.add_paragraph('Прибыль: {{profit}} руб.')
    doc.add_paragraph('')
    
    # Детализация
    doc.add_heading('Детализация', level=1)
    doc.add_paragraph('{{details}}')
    
    doc.save('templates/report_template.docx')
    print("Шаблон отчета создан: templates/report_template.docx")

# Шаблон справки
def create_certificate_template():
    doc = Document()
    
    # Заголовок
    title = doc.add_heading('СПРАВКА', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('')
    doc.add_paragraph('{{content}}')
    doc.add_paragraph('')
    doc.add_paragraph('')
    
    # Подпись
    doc.add_paragraph('{{signature}}')
    doc.add_paragraph('{{date}}')
    
    doc.save('templates/certificate_template.docx')
    print("Шаблон справки создан: templates/certificate_template.docx")

if __name__ == '__main__':
    create_contract_template()
    create_report_template()
    create_certificate_template()
    print("\nВсе шаблоны успешно созданы!")

